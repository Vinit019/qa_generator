import os
import json
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportUtils:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()
    
    def setup_custom_styles(self):
        """Setup custom styles for PDF export"""
        self.styles.add(ParagraphStyle(
            name='QuestionStyle',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            leftIndent=20
        ))
        
        self.styles.add(ParagraphStyle(
            name='AnswerStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=12,
            leftIndent=40,
            textColor=colors.grey
        ))
        
        self.styles.add(ParagraphStyle(
            name='HeaderStyle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            alignment=1  # Center alignment
        ))

def export_to_pdf(questions: dict, filename: str) -> str:
    """Export questions to PDF format"""
    export_utils = ExportUtils()
    
    # Create output filename
    base_name = filename.rsplit('.', 1)[0]
    output_path = f"outputs/questions_{base_name}.pdf"
    
    # Create PDF document
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = []
    
    # Add title
    title = Paragraph("Generated Questions", export_utils.styles['HeaderStyle'])
    story.append(title)
    story.append(Spacer(1, 20))
    
    # Add MCQ questions
    if questions.get('mcq'):
        story.append(Paragraph("Multiple Choice Questions (1 Mark Each)", export_utils.styles['Heading2']))
        story.append(Spacer(1, 10))
        
        for i, mcq in enumerate(questions['mcq'], 1):
            # Question
            q_text = f"Q{i}. {mcq['question']}"
            story.append(Paragraph(q_text, export_utils.styles['QuestionStyle']))
            
            # Options
            for j, option in enumerate(mcq['options'], 1):
                option_text = f"({chr(96+j)}) {option}"
                story.append(Paragraph(option_text, export_utils.styles['Normal']))
            
            story.append(Spacer(1, 10))
    
    # Add Short Answer questions
    if questions.get('short_answer'):
        story.append(Paragraph("Short Answer Questions (2 Marks Each)", export_utils.styles['Heading2']))
        story.append(Spacer(1, 10))
        
        for i, saq in enumerate(questions['short_answer'], 1):
            q_text = f"Q{i}. {saq['question']}"
            story.append(Paragraph(q_text, export_utils.styles['QuestionStyle']))
            story.append(Spacer(1, 5))
            
            # Sample answer
            answer_text = f"Sample Answer: {saq['sample_answer']}"
            story.append(Paragraph(answer_text, export_utils.styles['AnswerStyle']))
            story.append(Spacer(1, 10))
    
    # Add Long Answer questions
    if questions.get('long_answer'):
        story.append(Paragraph("Long Answer Questions (5 Marks Each)", export_utils.styles['Heading2']))
        story.append(Spacer(1, 10))
        
        for i, laq in enumerate(questions['long_answer'], 1):
            q_text = f"Q{i}. {laq['question']}"
            story.append(Paragraph(q_text, export_utils.styles['QuestionStyle']))
            story.append(Spacer(1, 5))
            
            # Detailed answer
            answer_text = f"Sample Answer: {laq['detailed_answer']}"
            story.append(Paragraph(answer_text, export_utils.styles['AnswerStyle']))
            story.append(Spacer(1, 15))
    
    # Build PDF
    doc.build(story)
    return output_path

def export_to_docx(questions: dict, filename: str) -> str:
    """Export questions to Word document format"""
    # Create output filename
    base_name = filename.rsplit('.', 1)[0]
    output_path = f"outputs/questions_{base_name}.docx"
    
    # Create Word document
    doc = DocxDocument()
    
    # Add title
    title = doc.add_heading('Generated Questions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add MCQ questions
    if questions.get('mcq'):
        doc.add_heading('Multiple Choice Questions (1 Mark Each)', level=1)
        
        for i, mcq in enumerate(questions['mcq'], 1):
            # Question
            q_text = f"Q{i}. {mcq['question']}"
            doc.add_paragraph(q_text)
            
            # Options
            for j, option in enumerate(mcq['options'], 1):
                option_text = f"({chr(96+j)}) {option}"
                doc.add_paragraph(option_text, style='List Bullet')
            
            doc.add_paragraph()  # Empty line
    
    # Add Short Answer questions
    if questions.get('short_answer'):
        doc.add_heading('Short Answer Questions (2 Marks Each)', level=1)
        
        for i, saq in enumerate(questions['short_answer'], 1):
            q_text = f"Q{i}. {saq['question']}"
            doc.add_paragraph(q_text)
            
            # Sample answer
            answer_text = f"Sample Answer: {saq['sample_answer']}"
            doc.add_paragraph(answer_text, style='Quote')
            doc.add_paragraph()  # Empty line
    
    # Add Long Answer questions
    if questions.get('long_answer'):
        doc.add_heading('Long Answer Questions (5 Marks Each)', level=1)
        
        for i, laq in enumerate(questions['long_answer'], 1):
            q_text = f"Q{i}. {laq['question']}"
            doc.add_paragraph(q_text)
            
            # Detailed answer
            answer_text = f"Sample Answer: {laq['detailed_answer']}"
            doc.add_paragraph(answer_text, style='Quote')
            doc.add_paragraph()  # Empty line
    
    # Save document
    doc.save(output_path)
    return output_path
