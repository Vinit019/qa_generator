import os
import PyPDF2
from docx import Document
import re
from typing import str

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from PDF or Word document
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        
        return self._clean_text(text)
    
    def _extract_from_word(self, file_path: str) -> str:
        """
        Extract text from Word document
        """
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and preprocess extracted text
        """
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-]', '', text)
        
        # Remove page numbers and headers/footers
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        return text.strip()
    
    def get_document_info(self, file_path: str) -> dict:
        """
        Get basic information about the document
        """
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        
        return {
            'name': file_name,
            'size': file_size,
            'format': os.path.splitext(file_path)[1].lower()
        }
